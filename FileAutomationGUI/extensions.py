from .logbase import logger

import tkinter as tk
from tkinter import ttk
# import table_logic
import docx
import PyPDF2
from .image_handling import ImageHandler
import subprocess, os, platform
from .audio import Audio


class App(tk.Tk):
    def __init__(self, data):
        logger.create_log(logger_name="extensions.py",
                          message="A new window titled 'Select the file extension', to set the merged file parameters has open.",
                          message_level="INFO")
        super().__init__()
        self.geometry("700x300")
        self.title('Select the file extension')

        # initialize data
        self.extension = ("None",'.txt', '.docx', ".pdf", ".mp3", ".png", ".jpg")

        # set up variable
        self.option_var = tk.StringVar(self)
        self.dst = tk.StringVar()
        self.file_paths = data
        self.file = None
        self.ext = None
        self.file_text = None
        self.output = None
        self.paddings = {'padx': 5, 'pady': 5}
        self.combined = None
        self.img_hndlr = ImageHandler()

        self.failed_files = 0
        self.passed_files = 0

        # create widget
        self.create_wigets()

    def create_wigets(self):
        # padding for widgets using the grid layout
        # label
        label = ttk.Label(self,  text='Save to the file with extension:')
        label.grid(column=0, row=0, sticky=tk.W, **self.paddings)

        # option menu
        option_menu = ttk.OptionMenu(
            self,
            self.option_var,
            self.extension[0],
            *self.extension,
            command=self.option_changed)

        option_menu.grid(column=1, row=0, sticky=tk.W, **self.paddings)

        # output label
        self.output_label = ttk.Label(self, foreground='red')
        self.output_label.grid(column=0, row=1, sticky=tk.W, **self.paddings)

        # File name

        label1 = ttk.Label(self, text='Save with file name:')
        label1.grid(column=0, row=3, sticky=tk.W, **self.paddings)

        self.file = tk.Entry(self, border=5, font = "Calibri 10", width=50, textvariable=self.dst)
        self.file.bind('<Return>', self.merge)

        self.file.grid(column=1, row=3)

    def option_changed(self, *args):
        logger.create_log(logger_name="table.py",
                          message=f"You selected the merged file extension as {self.option_var.get()}",
                          message_level="INFO")
        self.output_label['text'] = f'You selected: {self.option_var.get()}'
        self.ext = self.option_var.get()

    def merge(self, *args):
        self.passed_files = 0
        self.failed_files = 0

        if self.ext == ".docx":
            self.new_doc = docx.Document()

        if self.ext == ".pdf":
            self.new_pdf = PyPDF2.PdfFileWriter
            self.pdf_merger = PyPDF2.PdfFileMerger()



        logger.create_log(logger_name="table.py",
                          message=f"You entered merged file name as {self.file.get()} with {self.ext} extension.",
                          message_level="INFO")

        frame1 = tk.Frame(self)
        frame1.grid(column=1, row=5, sticky=tk.W)

        self.dst = self.file.get()

        try:
            if not self.dst.endswith(self.ext):
                self.dst = self.dst + self.ext
        except TypeError as e:
            logger.create_log(logger_name="table.py",
                              message=f"Below error has occured!",
                              message_level="ERROR")
            logger.create_log(logger_name="table.py",
                              message=f"Exception => {e}",
                              message_level="EXCEPTION")


        logger.create_log(logger_name="table.py",
                          message=f"Merging process for the file {self.file.get()} with {self.ext} extension has begun...",
                          message_level="INFO")

        logger.create_log(logger_name="table.py", message="Below files failed to merge!!",
                          message_level="WARNING")
        print(self.file_paths)
#####################################################################
        for src in self.file_paths:
            if self.ext == ".png" or self.ext == ".jpg":
                self.img_hndlr.merge_img(*self.file_paths, save_as=self.dst, extn=self.ext)
                # message
                self.passed_files = "Ignore"
                self.failed_files = "Ignore"
                if self.img_hndlr.wrong_files:
                    logger.create_log(logger_name="table.py",
                                      message=f"{self.img_hndlr.correct_files}/{self.img_hndlr.wrong_files + self.img_hndlr.correct_files} files merged!",
                                      message_level="WARNING")
                    self.error_log = tk.Label(frame1, text=f"{self.img_hndlr.correct_files}/{self.img_hndlr.wrong_files+self.img_hndlr.correct_files} files merged!\nCheck the log file for more details...")
                    self.error_log.grid()

                else:
                    logger.create_log(logger_name="table.py",
                                      message=f"All the files({self.img_hndlr.correct_files}/{self.img_hndlr.wrong_files + self.img_hndlr.correct_files}) merged successfully!",
                                      message_level="WARNING")
                    self.error_log = tk.Label(frame1,
                                              text=f"All the files({self.img_hndlr.correct_files}/{self.img_hndlr.wrong_files + self.img_hndlr.correct_files}) merged successfully!")
                    self.error_log.grid()

                break

            elif self.ext == ".mp3":
                aud = Audio(self.file_paths, dst=self.dst)
                aud.process()
                break


            else:
                self.write_to_file(src=src, dst=self.dst, ext=self.ext)


        if not self.failed_files and self.passed_files and self.failed_files!="Ignore":
            logger.create_log(logger_name="table.py",
                              message=f"All the files({self.passed_files}/{self.passed_files + self.failed_files}) merged successfully to {self.dst}!\nSee the log file for more details...",
                              message_level="WARNING")
            self.error_log = tk.Label(frame1,
                                      text=f"All the files({self.passed_files}/{self.passed_files + self.failed_files}) merged successfully to {self.dst}!\nSee the log file for more details...")
            self.error_log.grid()

        elif self.failed_files and self.failed_files!="Ignore":
            logger.create_log(logger_name="table.py",
                              message=f"Only ({self.passed_files}/{self.passed_files + self.failed_files}) files merged successfully to {self.dst}!\nSee the log file for more details...",
                              message_level="WARNING")
            self.error_log = tk.Label(frame1,
                                      text=f"Only ({self.passed_files}/{self.passed_files + self.failed_files}) files merged successfully to {self.dst}!\nSee the log file for more details...")
            self.error_log.grid()

        elif self.failed_files!="Ignore":
            logger.create_log(logger_name="table.py",
                              message=f"Merge to file {self.dst} failed!",
                              message_level="ERROR")
            self.error_log = tk.Label(frame1,
                                      text=f"Only ({self.passed_files}/{self.passed_files + self.failed_files}) files merged successfully to {self.dst}!\nSee the log file for more details...")
            self.error_log.grid()





        self.output = tk.Label(frame1, text = f"({self.dst}) created! Click here to view.")
        self.output.grid()

        tk.Button(frame1, text="Open", command=self.open_file).grid()


    def open_file(self):

        logger.create_log(logger_name="table.py",
                          message=f"You clicked the Open button to open the file {self.file.get()} with {self.ext} extension...",
                          message_level="INFO")

        filepath = self.dst

        if platform.system() == 'Darwin':  # macOS
            subprocess.call(('open', filepath))
        elif platform.system() == 'Windows':  # Windows
            os.startfile(filepath)
        else:  # linux variants
            subprocess.call(('xdg-open', filepath))




    def write_to_file(self, src, dst, ext):
        """
        This will write all the data of the files in the source location to a single destination file with the proper formatting. This function has the
        power to extract all the file with given extension (ext) even from the sub-directories. Thus, you can give the source location as the outermost directory.
        :param src:
        :param dst:
        :param ext:
        :return:
        """



        try:
            if ext == ".txt":
                self.combined = open(dst, "a+", encoding="utf-8")
                try:
                    file_read = open(src, "r", encoding="utf-8").read()
                    self.combined.writelines(f"\n\n----------------- {src} ------------------\n\n")
                    self.combined.write(file_read)

                    self.passed_files += 1
                    self.combined.close()

                except UnicodeDecodeError as e:
                    self.combined.close()
                    logger.create_log(logger_name="extensions.py", message_level="ERROR", message=f"ERROR!!: {src}")
                    logger.create_log(logger_name="extensions.py", message_level="EXCEPTION", message=e)

                    self.failed_files += 1

            elif ext == ".docx":
                try:
                    try:
                        doc = docx.Document(src)
                    except Exception as e:
                        logger.create_log(logger_name="extensions.py", message_level="ERROR", message=f"ERROR!!: {src}")
                        logger.create_log(logger_name="extensions.py", message_level="EXCEPTION", message=e)
                    else:
                        data = ""
                        # full_text = []

                        self.new_doc.add_heading(f"{src}", 3)
                        for para in doc.paragraphs:
                            data += para.text + "\n"

                        self.new_doc.add_paragraph(data)
                        self.new_doc.save(dst)

                        self.passed_files += 1

                except Exception as e:
                    logger.create_log(logger_name="extensions.py", message_level="ERROR", message=f"ERROR!!: {src}")
                    logger.create_log(logger_name="extensions.py", message_level="EXCEPTION", message=e)

                    self.failed_files += 1



            elif ext == ".pdf":
                try:
                    try:
                        self.pdf_merger.append(src)
                    except:
                        raise Exception


                    with open(dst, "wb") as output_file:
                        self.pdf_merger.write(output_file)

                    self.passed_files += 1
                except Exception as e:
                    logger.create_log(logger_name="extensions.py", message_level="ERROR", message=f"ERROR!!: {src}")
                    logger.create_log(logger_name="extensions.py", message_level="EXCEPTION", message=e)

                    # label_1 = ttk.Label(self, text=e, foreground="red")
                    # label_1.grid(column=0, row=5, sticky=tk.W, **self.paddings)

                    self.failed_files += 1

                    logger.create_log(logger_name="extensions.py", message_level="ERROR", message=f"ERROR!!: {src}")
                    logger.create_log(logger_name="extensions.py", message_level="EXCEPTION", message=e)

        except IOError as e:
            logger.create_log(logger_name="extensions.py", message_level="ERROR", message=f"ERROR!!: {src}")
            logger.create_log(logger_name="extensions.py", message_level="EXCEPTION", message=e)

            label2 = ttk.Label(self, text='There was an error opening the file', foreground="red")
            label2.grid(column=0, row=5, sticky=tk.W, **self.paddings)

    def get_instance(self):
        return self